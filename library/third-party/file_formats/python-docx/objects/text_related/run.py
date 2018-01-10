Run :::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

# 

# :::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

from docx import Document

# :::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

paragraph = Document(docx_file).paragraphs[0]
run = paragraph.runs[0]

# :::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

add_break(break_type=6)
#

add_picture(image_path_or_stream, width=None, height=None)
#

add_tab()
#

add_text(text)
#

bold
#

clear()
#

font
#

italic
#

style
#

text
#

underline
#