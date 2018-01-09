python-docx ::::::::::::::::::::::::::::::::::::::::::::::::::::

# python-docx - это библиотека Python для создания и обновления
# файлов Microsoft Word (.docx).

# документация на английском:
# https://python-docx.readthedocs.io/en/latest/

# установка:
pip install python-docx

# ::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

from docx import *

# ::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

doc = Document('demo.docx')

paragraph = doc.paragraphs[0]

run_obj = paragraph.runs[0]

core_properties = doc.core_properties

inline_shapes = doc.inline_shapes

sections = doc.sections

section = doc.sections[0]

settings = doc.settings

styles = run_obj.styles

base_style = styles.base_style

latent_styles = styles.latent

tables = doc.tables

table = doc.tables[0]

font = run_obj.font

color = run_obj.style.color

color_format = color.ColorFormat

tab_stop = paragraph.tabstops[0]


# TODO: дописать документацию
# TODO: прописать пути к документации

