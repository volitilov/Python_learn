from docx import Document

# ::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::


add_heading(text=u'', level=1)
#

add_page_break()
#

add_paragraph(text=u'', style=None)
#

add_picture(image_path_or_stream, width=None, height=None)
#

add_section(start_type=2)
#

add_table(rows, cols, style=None)
#

core_properties
#
	author
	#

	category
	#

	comments
	#

	content_status
	#

	created
	#

	identifier
	#

	keywords
	#

	language
	#

	last_modified_by
	#

	last_printed
	#

	modified
	#

	revision
	#

	subject
	#

	title
	#

	version
	#


inline_shapes
#

paragraphs
# 

part
#

save(path_or_stream)
#

sections
#

settings
#
	element
	#

styles
#
	add_style(name, style_type, builtin=False)
	#

	default(style_type)
	#

	element
	#

	latent_styles
	#

	builtin
	#

	delete()
	#

	hidden
	#

	locked
	#

	name
	#

	priority
	#

	quick_style
	#

	unhide_when_used
	#		

tables
#