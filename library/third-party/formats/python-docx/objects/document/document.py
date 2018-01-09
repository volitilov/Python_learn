Document :::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

#

# ::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

from docx import Document

# ::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::

doc = Document(docx_file)

# ::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::::


add_heading(text=u'', level=1)
#

add_page_break()
#

add_paragraph(text=u'', style=None)
#

add_picture(image_path_or_stream, width=None, height=None)
#

add_section(start_type=2)
#

add_table(rows, cols, style=None)
#

core_properties
#

inline_shapes
#

paragraphs
# 

part
#

save(path_or_stream)
#

sections
#

settings
#

styles
#

tables
#